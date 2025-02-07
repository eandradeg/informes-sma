import streamlit as st
import pandas as pd
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
import re, locale, os
import openpyxl
from io import BytesIO
from PIL import Image as PILImage
from docx.shared import Inches, Cm, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Configura la localización a español
locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8')

# ----------------- Funciones originales -----------------

def extract_info_from_filename(filename):
    """
    Extrae la información del nombre del archivo Word, en particular la parroquia.
    Se espera un formato: PARROQUIA_XXG_OPERADORA_TIPO.docx
    """
    match = re.match(r"(.*)_(\dG)_(.*)_(.*)\.docx", filename)
    if match:
        parroquia = match.group(1).replace("_", " ").upper()
        tecnologia = match.group(2)
        operadora = match.group(3).upper()
        tipo_medicion = match.group(4).upper()
        # Estandarizar el nombre del operador
        if operadora == "CONECEL":
            operadora = "CONECEL S.A."
        elif operadora == "OTECEL":
            operadora = "OTECEL S.A."
        st.write(f"[DEBUG] Archivo Word procesado: Parroquia={parroquia}, Tecnología={tecnologia}, Operadora={operadora}, Tipo={tipo_medicion}")
        return parroquia, tecnologia, operadora, tipo_medicion
    st.write("[DEBUG] Formato de nombre de archivo incorrecto.")
    return None, None, None, None

def replace_placeholder(paragraph, placeholder, new_text):
    """Reemplaza en un párrafo el marcador por el nuevo texto conservando el formato."""
    for run in paragraph.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, new_text)

def format_date(date_value, format_type):
    """Formatea la fecha según el formato solicitado."""
    try:
        date_obj = pd.to_datetime(date_value)
        if format_type == "month_only":
            return date_obj.strftime("%B")  # Ej: "enero"
        elif format_type == "dd_mm_yyyy":
            return date_obj.strftime("%d/%m/%Y")  # Ej: "31/01/2025"
        elif format_type == "long_format":
            return date_obj.strftime("%d de %B de %Y")  # Ej: "31 de enero de 2025"
        else:
            return str(date_value)
    except Exception as e:
        st.write(f"[DEBUG] Error en format_date: {e}")
        return str(date_value)

def replace_texts(paragraph, placeholders, counter):
    """
    Reemplaza marcadores en el párrafo.
    Usa un contador mutable para las fechas con el marcador «FECHA_CRONOGRAMA_DE_MEDICION_2024».
    """
    date_marker = "«FECHA_CRONOGRAMA_DE_MEDICION_2024»"
    if date_marker in paragraph.text:
        if counter[0] == 0:
            replace_placeholder(paragraph, date_marker, placeholders["fecha_antecedentes"])
        elif counter[0] == 1:
            replace_placeholder(paragraph, date_marker, placeholders["fecha_pruebas_realizadas"])
        elif counter[0] == 2:
            replace_placeholder(paragraph, date_marker, placeholders["fecha_conclusiones"])
        counter[0] += 1
    for key, value in placeholders.items():
        if key not in ["fecha_antecedentes", "fecha_pruebas_realizadas", "fecha_conclusiones"]:
            if key in paragraph.text:
                replace_placeholder(paragraph, key, value)
    return counter

def process_headers_and_footers(doc, placeholders):
    """Procesa encabezados y pies de página del documento Word,
    asegurando que los marcadores de fecha y número de informe se inserten como texto plano.
    """
    for section in doc.sections:
        header = section.header
        footer = section.footer

        # Procesar encabezados que estén dentro de tablas (si existen)
        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        # Procesar fecha de informe con "Cuenca,"
                        if "«FECHA_DE_INFORME»" in paragraph.text:
                            paragraph.clear()
                            run = paragraph.add_run("Cuenca, ")  # Agregar "Cuenca, "
                            run.font.name = 'Arial'
                            run.font.size = Pt(9)
                            run = paragraph.add_run(placeholders.get("«FECHA_DE_INFORME»", ""))
                            run.font.name = 'Arial'
                            run.font.size = Pt(9)
                        
                        # Procesar número de informe
                        if "«NÚMERO__DE_INFORME»" in paragraph.text:
                            paragraph.clear()
                            run = paragraph.add_run(placeholders.get("«NÚMERO__DE_INFORME»", ""))
                            run.font.name = 'Arial'
                            run.font.size = Pt(10)

        # Procesar párrafos del encabezado fuera de tablas
        for paragraph in header.paragraphs:
            # Procesar fecha de informe con "Cuenca,"
            if "«FECHA_DE_INFORME»" in paragraph.text:
                paragraph.clear()
                run = paragraph.add_run("Cuenca, ")  # Agregar "Cuenca, "
                run.font.name = 'Arial'
                run.font.size = Pt(9)
                run = paragraph.add_run(placeholders.get("«FECHA_DE_INFORME»", ""))
                run.font.name = 'Arial'
                run.font.size = Pt(9)
            
            # Procesar número de informe
            if "«NÚMERO__DE_INFORME»" in paragraph.text:
                paragraph.clear()
                run = paragraph.add_run(placeholders.get("«NÚMERO__DE_INFORME»", ""))
                run.font.name = 'Arial'
                run.font.size = Pt(10)

        # El procesamiento del pie de página permanece sin cambios
        for paragraph in footer.paragraphs:
            replace_texts(paragraph, placeholders, [0])
        for table in footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_texts(paragraph, placeholders, [0])

def process_doc_elements(doc, placeholders, selected_supervisor, recomendaciones):
    """
    Recorre los elementos (párrafos y tablas) del cuerpo del documento,
    reemplazando marcadores y actualizando según el supervisor seleccionado.
    """
    counter = [0]
    for element in list(doc.element.body):
        if element.tag.endswith('p'):
            paragraph = Paragraph(element, doc)
            replace_texts(paragraph, placeholders, counter)
        elif element.tag.endswith('tbl'):
            table = Table(element, doc)
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_texts(paragraph, placeholders, counter)
                        replace_placeholder(paragraph, "Ing. Mauricio Sánchez Pinos", selected_supervisor)
                        if selected_supervisor == "Ing. Mesías Vizuete López":
                            replace_placeholder(paragraph, "PROFESIONAL TÉCNICO 1", "ANALISTA TÉCNICO 2")
    
    # Reemplazar párrafos de recomendaciones
    if recomendaciones:  # Verifica si hay texto de recomendaciones para insertar
        start_recomendaciones = False  # Bandera para marcar cuando encontremos la sección
        recommendations_added = False   # Bandera para controlar que solo se inserte una vez
        for paragraph in doc.paragraphs:  # Itera sobre todos los párrafos del documento
            # Detecta el inicio de la sección de recomendaciones
            if "RECOMENDACIONES" in paragraph.text:
                start_recomendaciones = True
                continue  # Salta este párrafo para no modificar el título
            if start_recomendaciones:
                # Detecta el final de la sección de recomendaciones
                if "Informe realizado por:" in paragraph.text:
                    start_recomendaciones = False
                    break  # Sale del bucle al encontrar el final
                # Si aún no se han agregado las recomendaciones
                if not recommendations_added:
                    paragraph.clear()  # Limpia el contenido actual del párrafo
                    paragraph.add_run(recomendaciones)  # Agrega el nuevo texto
                    recommendations_added = True  # Marca que ya se agregaron las recomendaciones
                else:
                    paragraph.clear()  # Solo limpia los párrafos restantes
    return counter

# ----------------- Funciones nuevas para buscar e insertar el gráfico -----------------

def buscar_grafico(parroquia, carpeta):
    """
    Extrae imágenes PNG incrustadas directamente en la hoja de Excel.
    Ahora selecciona la SEGUNDA imagen en la hoja.
    """
    st.write(f"[DEBUG] Buscando imagen PNG para {parroquia} en {carpeta}")
    for archivo in os.listdir(carpeta):
        if archivo.startswith("~$") or not archivo.lower().endswith(".xlsx"):
            continue  # Solo procesar archivos .xlsx (evitar .xls y temporales)
        if parroquia.upper() in archivo.upper():
            ruta_excel = os.path.join(carpeta, archivo)
            st.write(f"[DEBUG] Excel encontrado: {ruta_excel}")
            # Intento 1: Extraer imagen usando openpyxl
            try:
                wb = openpyxl.load_workbook(ruta_excel, data_only=True)
                sheet = wb["MAPAS SMA-QoS-9"]
                # Verificar si hay imágenes en la hoja
                if not hasattr(sheet, "_images") or len(sheet._images) < 2:
                    st.error("No se encontraron suficientes imágenes en la hoja.")
                    return None, None
                # Tomar la SEGUNDA imagen
                imagen = sheet._images[1]  # Cambio aquí: seleccionamos la segunda imagen
                img_data = BytesIO(imagen._data())
                img_pil = PILImage.open(img_data)
                # Guardar temporalmente como PNG
                img_path = os.path.abspath(f"temp_{parroquia}_imagen.png")
                img_pil.save(img_path)
                st.write(f"[DEBUG] Imagen extraída y guardada en: {img_path}")
                # Mostrar la imagen en Streamlit
                st.image(img_path, caption="Imagen extraída de Excel", use_column_width=True)
                return img_path, ruta_excel
            except Exception as e:
                st.error(f"Error con openpyxl: {str(e)}")
                return None, None
    return None, None

def insertar_grafico_en_word(doc, img_path, ruta_excel):
    """
    Inserta el gráfico en el documento Word después del texto 'RESULTADOS'.
    """
    st.write(f"[DEBUG] Insertando gráfico en el documento")
    encontrado = False
    posicion = -1

    # Buscar el párrafo exacto donde aparece "RESULTADOS CONECEL S.A."
    for i, para in enumerate(doc.paragraphs):
        if "RESULTADOS CONECEL S.A." in para.text or "RESULTADOS OTECEL S.A." in para.text:
            st.write(f"[DEBUG] Sección encontrada en el párrafo {i}")
            posicion = i
            encontrado = True
            break

    if encontrado and posicion != -1:
        # Insertar la imagen directamente después del párrafo encontrado
        st.write(f"[DEBUG] Insertando gráfico en el párrafo {posicion + 1}")
        run = doc.paragraphs[posicion].add_run()
        run.add_picture(img_path, width=Inches(6))  # Ajustar el ancho a 6 pulgadas (más razonable)

        st.write("[DEBUG] Gráfico insertado correctamente.")
    else:
        st.error("[ERROR] No se encontró la sección 'RESULTADOS CONECEL S.A.' en el documento.")

def buscar_imagenes_encabezado_pie(carpeta):
    """
    Busca las imágenes PNG llamadas 'encabezado' y 'pie de página' en la carpeta especificada.
    """
    encabezado_path = None
    pie_path = None

    for archivo in os.listdir(carpeta):
        if archivo.lower() == "encabezado.png":
            encabezado_path = os.path.join(carpeta, archivo)
            st.write(f"[DEBUG] Imagen de encabezado encontrada: {encabezado_path}")
        elif archivo.lower() == "pie de pagina.png":
            pie_path = os.path.join(carpeta, archivo)
            st.write(f"[DEBUG] Imagen de pie de página encontrada: {pie_path}")

    return encabezado_path, pie_path

def reemplazar_imagenes_encabezado_pie(doc, encabezado_path, pie_path):
    """
    Reemplaza las imágenes en el encabezado y pie de página del documento Word.
    """
    st.write(f"[DEBUG] Reemplazando imágenes en encabezado y pie de página")

    for section in doc.sections:
        header = section.header
        footer = section.footer

        # Eliminar imágenes existentes en el encabezado
        for paragraph in header.paragraphs:
            for run in paragraph.runs:
                if run._element.xpath(".//w:drawing"):
                    run._element.clear()

        # Insertar nueva imagen en el encabezado
        header_paragraph = header.paragraphs[0]
        header_run = header_paragraph.add_run()
        header_run.add_picture(encabezado_path, width=Inches(6))

        # Eliminar imágenes existentes en el pie de página
        for paragraph in footer.paragraphs:
            for run in paragraph.runs:
                if run._element.xpath(".//w:drawing"):
                    run._element.clear()

        # Insertar nueva imagen en el pie de página
        footer_paragraph = footer.paragraphs[0]
        footer_run = footer_paragraph.add_run()
        footer_run.add_picture(pie_path, width=Inches(6))

    st.write("[DEBUG] Imágenes en encabezado y pie de página reemplazadas correctamente.")

def ajustar_margenes(doc):
    """
    Ajusta los márgenes del documento Word.
    """
    st.write(f"[DEBUG] Ajustando márgenes del documento")

    for section in doc.sections:
        section.top_margin = Cm(4)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    st.write("[DEBUG] Márgenes ajustados correctamente.")

def buscar_imagen_correccion_mapa(parroquia, carpeta):
    """
    Busca una imagen PNG o JPEG en la carpeta especificada que contenga el nombre de la parroquia.
    """
    st.write(f"[DEBUG] Buscando imagen de corrección de mapa para {parroquia} en {carpeta}")
    for archivo in os.listdir(carpeta):
        if archivo.lower().endswith((".png", ".jpeg", ".jpg")) and parroquia.upper() in archivo.upper():
            img_path = os.path.join(carpeta, archivo)
            st.write(f"[DEBUG] Imagen de corrección de mapa encontrada: {img_path}")
            return img_path
    return None

def insertar_imagen_correccion_mapa(doc, img_path):
    """
    Inserta la imagen de corrección de mapa en el documento Word después de 'Imagen 3.- Porcentaje de Cobertura WCDMA (3G), parámetro RSCP.'.
    """
    st.write(f"[DEBUG] Insertando imagen de corrección de mapa en el documento")
    encontrado = False
    posicion = -1

    # Buscar el párrafo exacto donde aparece "Imagen 3.- Porcentaje de Cobertura WCDMA (3G), parámetro RSCP."
    for i, para in enumerate(doc.paragraphs):
        if "Imagen 3.- Porcentaje de Cobertura WCDMA (3G), parámetro RSCP." in para.text:
            st.write(f"[DEBUG] Sección encontrada en el párrafo {i}")
            posicion = i
            encontrado = True
            break

    if encontrado and posicion != -1:
        # Insertar la imagen directamente después del párrafo encontrado
        st.write(f"[DEBUG] Insertando imagen de corrección de mapa en el párrafo {posicion + 1}")
        run = doc.paragraphs[posicion].add_run()
        run.add_picture(img_path, width=Inches(6))  # Ajustar el ancho a 6 pulgadas (más razonable)

        # Agregar el texto "Imagen 4.- Correción mapa de cobertura." con fuente Arial 9 y "Imagen.-" en negrita
        new_para = doc.add_paragraph()
        new_run = new_para.add_run("Imagen 4.- ")
        new_run.font.name = 'Arial'
        new_run.font.size = Pt(9)
        new_run.bold = True
        new_run = new_para.add_run("Corrección mapa de cobertura.")
        new_run.font.name = 'Arial'
        new_run.font.size = Pt(9)
        new_para.alignment = 1 # Centrar el texto

        st.write("[DEBUG] Imagen de corrección de mapa insertada correctamente.")
    else:
        st.error("[ERROR] No se encontró la sección 'Imagen 3.- Porcentaje de Cobertura WCDMA (3G), parámetro RSCP.' en el documento.")

# ----------------- Funciones para convertir Word a PDF -----------------

def convert_word_to_pdf(doc_path, pdf_path):
    """
    Convierte un documento de Word a PDF usando PyMuPDF.
    """
    doc = fitz.open(doc_path)
    pdf_bytes = doc.convert_to_pdf()
    with open(pdf_path, "wb") as f:
        f.write(pdf_bytes)
    st.write(f"[DEBUG] Documento Word convertido a PDF y guardado en: {pdf_path}")



# ----------------- Interfaz de Streamlit -----------------

st.title("Informes de Cobertura SMA")

# Carga de archivos para la parte de datos y para el documento Word
col1, col2 = st.columns(2)
with col1:
    uploaded_excel = st.file_uploader("Cargar archivo Consolidado Excel", type=["xls", "xlsx"])
with col2:
    uploaded_word = st.file_uploader("Cargar modelo de informe Word", type=["docx"])

# Botón de descarga al lado de la carga del informe de Word
if uploaded_word is not None:
    modified_filename = f"modified_{uploaded_word.name}"
    if os.path.exists(modified_filename):
        with open(modified_filename, "rb") as file:
            st.download_button(
                label="Descargar archivo Word modificado",
                data=file,
                file_name=modified_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

supervisors = [
    "Ing. Mauricio Sánchez Pinos",
    "Ing. Mesías Vizuete López",
    "Ing. Felipe Zumba Arichavala",
    "Ing. Ramiro Hurtado Figueroa"
]
selected_supervisor = st.selectbox("Selecciona al supervisor:", supervisors)

# Campo de texto para ingresar la ruta de la carpeta con los archivos Excel de gráficos
carpeta_graficos = st.text_input("Ingresa la ruta completa de la carpeta con archivos Excel y gráficos operadora",
                                  value=st.session_state.get("carpeta_graficos", ""))
st.session_state.carpeta_graficos = carpeta_graficos

# Campo de texto para ingresar la ruta de la carpeta con las imágenes de encabezado y pie de página
carpeta_imagenes = st.text_input("Ingresa la ruta completa de la carpeta con las imágenes de encabezado y pie de página",
                                  value=st.session_state.get("carpeta_imagenes", ""))
st.session_state.carpeta_imagenes = carpeta_imagenes

# Campo de texto para ingresar recomendaciones
recomendaciones = st.text_area("Ingresa recomendaciones (opcional):")

if uploaded_excel is not None and uploaded_word is not None:
    try:
        # Procesa el Excel de datos
        df_cobertura = pd.read_excel(uploaded_excel, sheet_name="COBERTURA")
        st.write("[DEBUG] Archivo Excel de datos cargado.")
        st.dataframe(df_cobertura)

        # Guarda el archivo Word cargado en un archivo temporal para manipularlo
        temp_word_path = f"temp_{uploaded_word.name}"
        with open(temp_word_path, "wb") as f:
            f.write(uploaded_word.getbuffer())
        st.write(f"[DEBUG] Archivo Word guardado temporalmente: {temp_word_path}")

        # Lee el documento Word
        doc = Document(temp_word_path)
        filename = uploaded_word.name
        st.write(f"[DEBUG] Archivo Word cargado: {filename}")

        # Extrae la información del nombre del archivo Word (incluida la parroquia)
        parroquia, tecnologia, operadora, tipo_medicion = extract_info_from_filename(filename)
        if parroquia and operadora:
            st.write(f"[DEBUG] Información extraída: Parroquia={parroquia}, Tecnología={tecnologia}, Operadora={operadora}, Tipo de Medición={tipo_medicion}")

            # Filtra el DataFrame por PARROQUIA y OPERADORA
            filtered_df = df_cobertura[
                (df_cobertura['PARROQUIA'].str.strip().str.upper() == parroquia) &
                (df_cobertura['OPERADORA'].str.strip() == operadora)
            ]
            if not filtered_df.empty:
                st.write("[DEBUG] DataFrame filtrado:")
                st.dataframe(filtered_df)

                # Extrae valores para los marcadores
                provincia = str(filtered_df['PROVINCIA'].values[0])
                canton = str(filtered_df['CANTÓN'].values[0])
                parroquia_val = str(filtered_df['PARROQUIA'].values[0])
                fecha_medicion = filtered_df['FECHA CRONOGRAMA DE MEDICION 2024'].values[0]
                numero_informe = str(filtered_df['NÚMERO  DE INFORME'].values[0])
                fecha_informe = filtered_df['FECHA DE INFORME'].values[0]
                numero_total_de_muestras_arcotel = str(filtered_df['NUMERO TOTAL DE MUESTRAS ARCOTEL'].values[0])
                numero_muestras_validas_arcotel = str(filtered_df['NUMERO VALIDAS ARCOTEL'].values[0])
                muestras_validas_velocidad = str(filtered_df['MUESTRAS VALIDAS VELOCIDAD ARCOTEL'].values[0])
                modificar_mapa_de_cobertura_arc = str(filtered_df['REQUIERE MODIFICAR MAPA DE COBERTURA ARCOTEL'].values[0])
                valor_medido = str(filtered_df['VALOR MEDIDO'].values[0])
                cobertura_operadora = str(filtered_df['COBERTURA OPERADORA'].values[0])
                alcanza_valor_objetivo_arcotel = str(filtered_df['ALCANZA VALOR OBJETIVO ARCOTEL'].values[0])
                muestras_validas_operadora = str(filtered_df['PORCENTAJE DE MUESTRAS VALIDAS OPERADORA'].values[0])
                alcanza_valor_objetivo_operadora = str(filtered_df['ALCANZA VALOR OBJETIVO OPERADORA'].values[0])
                modificar_mapa_de_cobertura_operadora = str(filtered_df['REQUIERE MODIFICAR MAPA DE COBERTURA OPERADORA'].values[0])

                placeholders = {
                    "fecha_antecedentes": format_date(fecha_medicion, "month_only"),
                    "fecha_pruebas_realizadas": format_date(fecha_medicion, "dd_mm_yyyy"),
                    "fecha_conclusiones": format_date(fecha_medicion, "long_format"),
                    "«PROVINCIA»": provincia,
                    "«Provincia»": provincia.capitalize(),
                    "«CANTÓN»": canton,
                    "«Cantón»": canton.capitalize(),
                    "«PARROQUIA»": parroquia_val,
                    "«Parroquia»": parroquia_val.capitalize(),
                    "«NÚMERO__DE_INFORME»": numero_informe,
                    "«FECHA_DE_INFORME»": format_date(fecha_informe, "long_format"),
                    "«VALOR_MEDIDO»": valor_medido,
                    "«COBERTURA_OPERADORA»": cobertura_operadora,
                    "«ALCANZA_VALOR_OBJETIVO_ARCOTEL»": alcanza_valor_objetivo_arcotel,
                    "«NUMERO_TOTAL_DE_MUESTRAS_ARCOTEL»": numero_total_de_muestras_arcotel,
                    "«NUMERO_VALIDAS_ARCOTEL»": numero_muestras_validas_arcotel,
                    "«MUESTRAS_VALIDAS_VELOCIDAD_ARCOTEL»": muestras_validas_velocidad,
                    "«REQUIERE_MODIFICAR_MAPA_DE_COBERTURA_ARC»": modificar_mapa_de_cobertura_arc,
                    "«PORCENTAJE_DE_MUESTRAS_VALIDAS_OPERADORA»": muestras_validas_operadora,
                    "«ALCANZA_VALOR_OBJETIVO_OPERADORA»": alcanza_valor_objetivo_operadora,
                    "«REQUIERE_MODIFICAR_MAPA_DE_COBERTURA_OPE»": modificar_mapa_de_cobertura_operadora
                }

                # Procesa el documento: reemplaza marcadores en el cuerpo, encabezados y pies
                process_doc_elements(doc, placeholders, selected_supervisor, recomendaciones)
                process_headers_and_footers(doc, placeholders)

                # Busca el archivo Excel que contenga la parroquia para extraer el gráfico
                ruta_carpeta = st.session_state.get("carpeta_graficos", "").strip()
                if ruta_carpeta != "":
                    st.write(f"[DEBUG] Ruta de la carpeta de gráficos: {ruta_carpeta}")
                    img_path, ruta_excel_graf = buscar_grafico(parroquia, ruta_carpeta)
                    st.write(f"[DEBUG] Ruta Excel del gráfico encontrado: {ruta_excel_graf}")
                    if img_path:
                        # Insertar la imagen directamente en el mismo documento
                        insertar_grafico_en_word(doc, img_path, ruta_excel_graf)

                        # Eliminar la imagen temporal después de insertarla
                        os.remove(img_path)
                        st.success("Gráfico insertado en la sección 'RESULTADOS'")
                    else:
                        st.warning("No se encontró un archivo Excel con la parroquia en el nombre o el gráfico no pudo extraerse.")
                else:
                    st.write("[DEBUG] La variable 'carpeta_graficos' está vacía. Ingrese una ruta válida en el campo de texto.")

                # Buscar y reemplazar imágenes en encabezado y pie de página
                ruta_carpeta_imagenes = st.session_state.get("carpeta_imagenes", "").strip()
                if ruta_carpeta_imagenes != "":
                    st.write(f"[DEBUG] Ruta de la carpeta de imágenes: {ruta_carpeta_imagenes}")
                    encabezado_path, pie_path = buscar_imagenes_encabezado_pie(ruta_carpeta_imagenes)
                    if encabezado_path and pie_path:
                        reemplazar_imagenes_encabezado_pie(doc, encabezado_path, pie_path)

                        # Ajustar márgenes del documento
                        ajustar_margenes(doc)

                # Si modificar_mapa_de_cobertura_operadora es "SI", insertar la imagen de corrección de mapa
                if modificar_mapa_de_cobertura_operadora.upper() == "SI":
                    img_correccion_path = buscar_imagen_correccion_mapa(parroquia, ruta_carpeta)
                    if img_correccion_path:
                        insertar_imagen_correccion_mapa(doc, img_correccion_path)

                # Guardar el archivo modificado final
                modified_filename = f"{numero_informe}_{filename}"
                doc.save(modified_filename)
                st.write(f"[DEBUG] Documento modificado guardado: {modified_filename}")

                # Ofrecer el archivo modificado para su descarga
                with open(modified_filename, "rb") as file:
                    st.download_button(
                        label="Descargar archivo Word modificado",
                        data=file,
                        file_name=modified_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

            else:
                st.error(f"No se encontraron datos para la parroquia {parroquia} con tecnología {tecnologia} y operadora {operadora}.")
        else:
            st.error("El nombre del archivo no sigue el formato esperado.")
    except Exception as e:
        st.error(f"Ocurrió un error: {e}")
else:
    st.info("Esperando a que cargues ambos archivos...")
